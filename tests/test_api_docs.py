import io
import pathlib
import sys

import pytest
from docx import Document
from fastapi.testclient import TestClient

ROOT = pathlib.Path(__file__).resolve().parents[1]
sys.path.append(str(ROOT))
sys.path.append(str(ROOT / "backend"))

from backend.app import DOCS, STATE_FILE, STORAGE, app


def _cleanup_storage() -> None:
    DOCS.clear()
    for path in STORAGE.glob("*"):
        if path.is_file():
            path.unlink(missing_ok=True)
    STATE_FILE.unlink(missing_ok=True)


def _create_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Lesoverzicht voor test")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Week"
    table.cell(0, 1).text = "Onderwerp"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Introductie"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture(autouse=True)
def reset_storage():
    _cleanup_storage()
    yield
    _cleanup_storage()


client = TestClient(app)


def test_docs_flow():
    payload = _create_docx_bytes()
    files = {
        "file": (
            "voorbeeld.docx",
            payload,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    upload_res = client.post("/api/uploads", files=files)
    assert upload_res.status_code == 200
    meta = upload_res.json()
    file_id = meta["fileId"]

    preview_res = client.get(f"/api/docs/{file_id}/preview")
    assert preview_res.status_code == 200
    preview = preview_res.json()

    assert preview["mediaType"].startswith("text/html")
    assert "Lesoverzicht" in preview["html"]
    assert "url" not in preview

    source_res = client.get(f"/api/docs/{file_id}/source")
    assert source_res.status_code == 200
    disposition = source_res.headers.get("content-disposition", "")
    assert "attachment" in disposition.lower()
