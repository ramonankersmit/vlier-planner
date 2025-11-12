from pathlib import Path
from typing import Iterable

import pytest

from backend.parsers.parser_docx import (
    extract_meta_from_docx,
    extract_rows_from_docx,
    extract_all_periods_from_docx,
)


def _make_period_2_sample(tmp_path: Path, weeks: Iterable[int]) -> Path:
    from docx import Document

    doc = Document()
    # Plaats een misleidende verwijzing naar periode 1 in de hoofdtekst.
    doc.add_paragraph("Terugblik Periode 1")
    # De daadwerkelijke periode 2 staat verderop, waardoor de eerste match verkeerd kan zijn.
    doc.add_paragraph("Planning Periode 2")

    table_weeks = list(weeks)
    table = doc.add_table(rows=len(table_weeks) + 1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Week"
    table.cell(0, 1).text = "Onderwerp"
    for idx, week in enumerate(table_weeks, start=1):
        table.cell(idx, 0).text = str(week)
        table.cell(idx, 1).text = f"Les {week}"

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "4 vwo · Periode 2"

    sample_path = tmp_path / "Maatschappijleer_4vwo_p2.docx"
    doc.save(sample_path)
    return sample_path


def test_latijnse_taal_en_cultuur_period_1_stops_before_period_2():
    sample = Path("samples/Latijnse Taal en Cultuur_ periode 1_4vwo.docx")
    assert sample.exists(), "Sample document is missing"

    meta = extract_meta_from_docx(str(sample), sample.name)
    assert meta.beginWeek == 35
    assert meta.eindWeek == 45

    weeks = [row.week for row in extract_rows_from_docx(str(sample), sample.name)]
    assert weeks, "Expected at least one extracted row"
    assert max(weeks) == 45
    assert all(35 <= week <= 45 for week in weeks)


def test_latijnse_taal_en_cultuur_period_2_wraps_across_new_year():
    sample = Path("samples/Latijnse Taal en Cultuur_ periode 1_4vwo.docx")
    assert sample.exists(), "Sample document is missing"

    rows = extract_rows_from_docx(str(sample), sample.name, target_periode=2)
    weeks = [row.week for row in rows]
    assert weeks, "Expected period 2 rows to be extracted"
    assert set(range(46, 53)).issubset(weeks)
    assert set(range(1, 5)).issubset(weeks)

    meta = extract_meta_from_docx(str(sample), sample.name, target_periode=2)
    assert meta.beginWeek == 46
    assert meta.eindWeek == 4


def test_latijnse_taal_en_cultuur_all_periods_are_detected():
    sample = Path("samples/Latijnse Taal en Cultuur_ periode 1_4vwo.docx")
    assert sample.exists(), "Sample document is missing"

    parsed = extract_all_periods_from_docx(str(sample), sample.name)
    assert len(parsed) == 2

    periods = [meta.periode for meta, _ in parsed]
    assert periods == [1, 2]

    first_meta, first_rows = parsed[0]
    assert first_meta.beginWeek == 35
    assert first_meta.eindWeek == 45
    first_weeks = [row.week for row in first_rows]
    assert first_weeks and max(first_weeks) == 45

    second_meta, second_rows = parsed[1]
    assert second_meta.beginWeek == 46
    assert second_meta.eindWeek == 4
    second_weeks = [row.week for row in second_rows]
    assert set(range(46, 53)).issubset(second_weeks)
    assert set(range(1, 5)).issubset(second_weeks)

@pytest.mark.parametrize(
    "filename",
    [
        "Natuurkunde_4vwo_P1.docx",
        "Maatschappijleer_4vwo_p1.docx",
        "Geschiedenis 4V P1 studiewijzer.docx",
        "Levensbeschouwing 2526 4V periode 1.docx",
        "Nederlands 4V planner periode 1 2526.docx",
    ],
)
def test_schooljaar_is_detected_for_samples(filename: str) -> None:
    sample = Path("samples") / filename
    assert sample.exists(), "Sample document is missing"

    meta = extract_meta_from_docx(str(sample), sample.name)
    assert meta.schooljaar == "2025/2026"


def test_periode_detections_use_footer_and_filename(tmp_path: Path) -> None:
    weeks = list(range(46, 53)) + list(range(1, 5))
    sample = _make_period_2_sample(tmp_path, weeks)

    meta = extract_meta_from_docx(str(sample), "Aardrijkskunde_4V_P2_2025-2026.docx")
    assert meta.periode == 2
    assert meta.beginWeek == 46
    assert meta.eindWeek == 4

    rows = extract_rows_from_docx(str(sample), "Aardrijkskunde_4V_P2_2025-2026.docx")
    extracted_weeks = [row.week for row in rows]
    assert set(range(46, 53)).issubset(extracted_weeks)
    assert set(range(1, 5)).issubset(extracted_weeks)
