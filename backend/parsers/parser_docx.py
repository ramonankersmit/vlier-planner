from collections import defaultdict
from dataclasses import dataclass
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import Optional, List, Tuple, Iterable, Dict, Set
import re

try:  # pragma: no cover - prefer package-relative imports when available
    from ..models import DocMeta, DocRow
except ImportError:  # pragma: no cover
    from models import DocMeta, DocRow  # type: ignore

from .base_parser import (
    BaseParser,
    RawEntry,
    WeekCellParseResult,
    extract_schooljaar_from_text,
)
from .config import get_keyword_config

# ---------------------------
# Regex patronen & helpers
# ---------------------------

RE_STUDIEWIJZER = re.compile(r"studiewijzer", re.I)
RE_VAK_IN_BRACKETS = re.compile(r"\[\s*([A-Za-zÀ-ÿ\s\-\&]+)\s*\]")
RE_VAK_AFTER_DASH = re.compile(r"studiewijzer\s*[-–]\s*(.+)", re.I)
RE_ANY_BRACKET_VAK = re.compile(r"\[\s*([A-Za-zÀ-ÿ\s\-\&]+?)\s*\]")
RE_PERIODE_MARKER = re.compile(r"periode\s*([1-4])", re.I)

KEYWORDS = get_keyword_config()
BASE_PARSER = BaseParser(KEYWORDS)

WEEK_HEADER_KEYWORDS = KEYWORDS.week_headers
DATE_HEADER_KEYWORDS = KEYWORDS.date_headers
LES_HEADER_KEYWORDS = KEYWORDS.lesson_headers
ONDERWERP_HEADERS = KEYWORDS.subject_headers
LEERDOEL_HEADERS = KEYWORDS.objective_headers
HUISWERK_HEADERS = KEYWORDS.homework_headers
OPDRACHT_HEADERS = KEYWORDS.assignment_headers
INLEVER_HEADERS = KEYWORDS.handin_headers
TOETS_HEADERS = KEYWORDS.exam_headers
BRON_HEADERS = KEYWORDS.resource_headers
NOTITIE_HEADERS = KEYWORDS.note_headers
KLAS_HEADERS = KEYWORDS.class_headers
LOCATIE_HEADERS = KEYWORDS.location_headers


def _clean(s: str) -> str:
    return (s or "").strip()


def _detect_primary_periode(doc: Document) -> Optional[int]:
    """Vind de eerste expliciete periodevermelding in de paragraaftekst."""
    for p in doc.paragraphs:
        txt = _clean(p.text)
        if not txt:
            continue
        m = RE_PERIODE_MARKER.search(txt)
        if m:
            try:
                return int(m.group(1))
            except ValueError:
                continue
    return None


def _parse_periode_from_filename(filename: str) -> Optional[int]:
    base = filename.rsplit(".", 1)[0]
    base = base.replace("_", " ").replace("-", " ")
    m = re.search(r"(?i)\bp\s*([1-4])\b", base)
    if m:
        try:
            return int(m.group(1))
        except ValueError:
            return None
    m = re.search(r"(?i)periode\s*([1-4])", base)
    if m:
        try:
            return int(m.group(1))
        except ValueError:
            return None
    return None


def _table_period_from_cells(tbl) -> Optional[int]:
    """Zoek naar een periodevermelding in de eerste rijen/kolommen van een tabel."""
    try:
        # Controleer de eerste paar rijen; veel bestanden zetten de periode
        # in de eerste rij of in een samengevoegde kopcel.
        for row in tbl.rows[:3]:
            for cell in row.cells[:3]:
                txt = _clean(cell.text)
                if not txt:
                    continue
                m = RE_PERIODE_MARKER.search(txt)
                if m:
                    try:
                        return int(m.group(1))
                    except ValueError:
                        continue
    except Exception:
        pass
    return None


def _table_period_markers(doc: Document) -> List[Tuple[Table, Optional[int]]]:
    """Geef een lijst terug met (tabel, periode-marker)."""
    results: List[Tuple[Table, Optional[int]]] = []
    current: Optional[int] = None
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            txt = _clean(paragraph.text)
            if not txt:
                continue
            m = RE_PERIODE_MARKER.search(txt)
            if m:
                try:
                    current = int(m.group(1))
                except ValueError:
                    current = None
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            marker = current or _table_period_from_cells(table)
            if marker is not None:
                current = marker
            results.append((table, marker))
    return results


def _table_matches_period(marker: Optional[int], periode: Optional[int]) -> bool:
    if marker is None or periode is None:
        return True
    return marker == periode

def _first_nonempty(paragraphs, start_idx: int) -> Optional[str]:
    for p in paragraphs[start_idx:]:
        t = _clean(p.text)
        if t:
            return t
    return None

# ---------------------------
# VAK parsing
# ---------------------------

def _parse_vak_anywhere(doc: Document) -> Optional[str]:
    limit = min(25, len(doc.paragraphs))
    for i in range(limit):
        txt = _clean(doc.paragraphs[i].text or "")
        m = RE_ANY_BRACKET_VAK.search(txt)
        if m:
            return _clean(m.group(1))
    return None

def vak_from_filename(filename: str) -> Optional[str]:
    base = filename.rsplit(".", 1)[0]
    base = base.replace("_", " ").replace("-", " ")
    base = re.sub(r"(?i)studiewijzer|planner|periode", " ", base)
    base = re.sub(r"(?i)\bp\s*\d+\b", " ", base)
    base = re.sub(r"\d+", " ", base)
    base = re.sub(r"(?i)\b(havo|vwo)\b", " ", base)
    tokens = [t for t in base.split() if len(t) > 1]
    cleaned = " ".join(tokens).strip()
    return cleaned or None

def _parse_vak_from_header(doc: Document, filename: str) -> Optional[str]:
    anywhere = _parse_vak_anywhere(doc)
    if anywhere:
        return anywhere

    for i, p in enumerate(doc.paragraphs):
        txt = _clean(p.text)
        if not txt:
            continue
        if RE_STUDIEWIJZER.search(txt):
            m = RE_VAK_AFTER_DASH.search(txt)
            if m:
                return _clean(m.group(1))
            nxt = _first_nonempty(doc.paragraphs, i + 1) or ""
            mb = RE_VAK_IN_BRACKETS.search(nxt)
            if mb:
                return _clean(mb.group(1))
            if i + 2 < len(doc.paragraphs):
                nxt2 = _clean(doc.paragraphs[i + 2].text or "")
                mb2 = RE_VAK_IN_BRACKETS.search(nxt2)
                if mb2:
                    return _clean(mb2.group(1))
    return vak_from_filename(filename)

# ---------------------------
# Schooljaar parsing
# ---------------------------

def _parse_schooljaar_from_filename(filename: str) -> Optional[str]:
    base = filename.rsplit(".", 1)[0]
    return extract_schooljaar_from_text(base)

def _parse_schooljaar_from_doc(doc: Document) -> Optional[str]:
    texts: List[str] = []
    for p in doc.paragraphs:
        texts.append(_clean(p.text))
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(_clean(cell.text))
    except Exception:
        pass
    try:
        sec = doc.sections[0]
        texts += [_clean(p.text) for p in sec.header.paragraphs]
        texts += [_clean(p.text) for p in sec.footer.paragraphs]
    except Exception:
        pass
    blob = " | ".join(t for t in texts if t)
    return extract_schooljaar_from_text(blob)

def _parse_footer_meta(doc: Document) -> Tuple[str, str, Optional[int], Optional[str]]:
    """Heuristiek uit footer: niveau (HAVO/VWO), leerjaar (1-6), periode (1-4), schooljaar (YYYY/YYYY, indien aanwezig)."""
    niveau = "VWO"
    leerjaar = "4"
    periode: Optional[int] = None
    schooljaar = None
    try:
        sec = doc.sections[0]
        footer_texts = [_clean(p.text) for p in sec.footer.paragraphs if _clean(p.text)]
        full = " | ".join(footer_texts)
        low = full.lower()
        if "havo" in low:
            niveau = "HAVO"
        if "vwo" in low:
            niveau = "VWO"
        m = re.search(r"\[\s*([1-6])\s*(havo|vwo)\s*\]", low)
        if m:
            leerjaar = m.group(1)
        else:
            m = re.search(r"\b([1-6])\b", low)
            if m:
                leerjaar = m.group(1)
        m = re.search(r"periode\s*([1-4])", low)
        if m:
            periode = int(m.group(1))
        candidate = extract_schooljaar_from_text(full)
        if candidate:
            schooljaar = candidate
    except Exception:
        pass
    return niveau, leerjaar, periode, schooljaar

# ---------------------------
# Weekrange parsing — eenvoudige regel
# ---------------------------

WEEK_HEADER_KEYWORDS = ("weeknummer", "week nr", "weeknr", "week", "wk")

def _find_week_column(headers: List[str]) -> Optional[int]:
    """Pak de eerste kolom waarvan de header 'week' of 'wk' bevat (case-insensitive, spaties genegeerd)."""
    norm = [" ".join(_clean(h).lower().split()) for h in headers]
    for i, h in enumerate(norm):
        for kw in WEEK_HEADER_KEYWORDS:
            nkw = " ".join(kw.lower().split())
            if nkw in h:
                return i
    return None

def _is_new_period(
    prev_week: Optional[int], current_week: int, *, allow_wrap: bool = False
) -> bool:
    """Detecteer overgang naar een nieuwe periode op basis van weeknummers."""
    if prev_week is None:
        return False
    if not (1 <= current_week <= 53):
        return False
    # Zodra de reeks na week 40+ terugvalt naar het begin van het jaar,
    # interpreteren we dat als een nieuwe periode.
    if allow_wrap:
        return False
    return prev_week >= 40 and current_week <= 10

def _table_rows_texts(tbl) -> List[List[str]]:
    """Converteer een docx-table naar matrix van celteksten, robuust genoeg voor merges."""
    rows: List[List[str]] = []
    for row in tbl.rows:
        rows.append([_clean(c.text) for c in row.cells])
    return rows


def _row_contains_weeks(row: List[str]) -> bool:
    for cell in row:
        if cell and parse_week_cell(cell):
            return True
    return False


def _combine_header_rows(header_rows: List[List[str]]) -> List[str]:
    if not header_rows:
        return []
    max_cols = max(len(row) for row in header_rows)
    combined: List[str] = []
    for col_idx in range(max_cols):
        parts: List[str] = []
        for row in header_rows:
            if col_idx >= len(row):
                continue
            text = normalize_text(row[col_idx] or "")
            if text:
                parts.append(text)
        combined.append(" ".join(parts))
    return combined


def _split_header_and_data_rows(rows: List[List[str]]) -> Tuple[List[str], List[List[str]], int]:
    if not rows:
        return [], [], 0

    header_rows: List[List[str]] = []
    data_start: Optional[int] = None
    for idx, row in enumerate(rows):
        if idx == 0:
            header_rows.append(row)
            continue
        if _row_contains_weeks([cell or "" for cell in row]):
            data_start = idx
            break
        header_rows.append(row)

    if data_start is None:
        data_start = 1 if len(rows) > 1 else len(rows)

    data_rows = rows[data_start:]
    if not header_rows:
        header_rows = [rows[0]]

    headers = _combine_header_rows(header_rows)
    if not any(normalize_text(h) for h in headers):
        headers = [normalize_text(c or "") for c in rows[0]]

    return headers, data_rows, len(header_rows)

def _parse_week_range(
    doc: Document, periode: Optional[int], table_markers: List[Tuple[Table, Optional[int]]]
) -> Tuple[int, int]:
    """
    - Neem in *elke* tabel rij 0 als header.
    - Vind de kolom waarvan de header 'week' of 'wk' bevat.
    - Verzamel ALLE weeknummers uit die kolom (over alle tabellen).
    - Neem min/max.
    """
    begin_w, eind_w = 0, 0
    ordered_weeks: List[int] = []
    unique_weeks: List[int] = []
    seen: set[int] = set()

    prev_week: Optional[int] = None
    stop = False

    try:
        for tbl, marker in table_markers:
            if not _table_matches_period(marker, periode):
                continue
            if stop:
                break
            rows = _table_rows_texts(tbl)
            if len(rows) < 2:
                continue
            headers = rows[0]
            week_col = _find_week_column(headers)
            if week_col is None:
                continue

            allow_wrap = (marker is not None) or (periode == 2)
            for r in rows[1:]:
                if stop:
                    break
                if week_col < len(r):
                    ws = _weeks_from_week_cell(r[week_col])
                    if not ws:
                        continue
                    for w in ws:
                        if not (1 <= w <= 53):
                            continue
                        if _is_new_period(prev_week, w, allow_wrap=allow_wrap):
                            stop = True
                            break
                        ordered_weeks.append(w)
                        if w not in seen:
                            unique_weeks.append(w)
                            seen.add(w)
                        prev_week = w

        if ordered_weeks:
            begin_w, eind_w = ordered_weeks[0], ordered_weeks[-1]
        elif unique_weeks:
            begin_w, eind_w = unique_weeks[0], unique_weeks[-1]

    except Exception:
        pass

    return begin_w, eind_w

# ---------------------------
# Hoofdfunctie
# ---------------------------

@dataclass
class _DocParseContext:
    doc: Document
    filename: str
    table_markers: List[Tuple[Table, Optional[int]]]
    niveau: Optional[str]
    leerjaar: Optional[str]
    periode_footer: Optional[int]
    periode_filename: Optional[int]
    schooljaar_footer: Optional[str]
    periode_text: Optional[int]
    schooljaar: Optional[str]
    vak: str


def _build_doc_context(path: str, filename: str) -> _DocParseContext:
    doc = Document(path)
    table_markers = _table_period_markers(doc)
    niveau, leerjaar, periode_footer, schooljaar_footer = _parse_footer_meta(doc)
    periode_text = _detect_primary_periode(doc)
    periode_filename = _parse_periode_from_filename(filename)
    schooljaar = (
        schooljaar_footer
        or _parse_schooljaar_from_doc(doc)
        or _parse_schooljaar_from_filename(filename)
    )
    vak = _parse_vak_from_header(doc, filename) or "Onbekend"
    return _DocParseContext(
        doc=doc,
        filename=filename,
        table_markers=table_markers,
        niveau=niveau,
        leerjaar=leerjaar,
        periode_footer=periode_footer,
        periode_filename=periode_filename,
        schooljaar_footer=schooljaar_footer,
        periode_text=periode_text,
        schooljaar=schooljaar,
        vak=vak,
    )


def _extract_meta_from_context(
    ctx: _DocParseContext, target_periode: Optional[int] = None
) -> Optional[DocMeta]:
    periode = (
        target_periode
        or ctx.periode_footer
        or ctx.periode_filename
        or ctx.periode_text
    )
    begin_week, eind_week = _parse_week_range(ctx.doc, periode, ctx.table_markers)
    file_id = re.sub(r"[^a-zA-Z0-9]+", "-", ctx.filename)[:40]

    final_periode = (
        periode
        or ctx.periode_footer
        or ctx.periode_filename
        or ctx.periode_text
        or 1
    )

    return DocMeta(
        fileId=file_id,
        bestand=ctx.filename,
        vak=ctx.vak,
        niveau=ctx.niveau,
        leerjaar=ctx.leerjaar,
        periode=final_periode,
        beginWeek=begin_week,
        eindWeek=eind_week,
        schooljaar=ctx.schooljaar,
    )


def extract_meta_from_docx(
    path: str, filename: str, target_periode: Optional[int] = None
) -> Optional[DocMeta]:
    ctx = _build_doc_context(path, filename)
    return _extract_meta_from_context(ctx, target_periode)


# -------------------------------------------------
# Nieuwe API: rij-niveau parsing
# -------------------------------------------------

from datetime import date

# Header keyword sets
DATE_HEADER_KEYWORDS = ("datum", "weekdatum", "date", "start", "begin")
LES_HEADER_KEYWORDS = ("les", "lesnr", "lesnummer")
normalize_text = BASE_PARSER.normalize_text
split_bullets = BASE_PARSER.split_bullets
find_header_idx = BASE_PARSER.find_header_idx
parse_week_cell = BASE_PARSER.parse_week_cell
parse_week_cell_details = BASE_PARSER.parse_week_cell_details
parse_date_cell = BASE_PARSER.parse_date_cell
parse_date_range_cell = BASE_PARSER.parse_date_range_cell
find_urls = BASE_PARSER.find_urls
parse_toets_cell = BASE_PARSER.parse_toets_cell
vak_from_filename = BASE_PARSER.vak_from_filename
_make_week_result = BaseParser._make_week_result


def _weeks_from_week_cell(text: str) -> List[int]:
    return BASE_PARSER.parse_week_cell(text)


def _note_contains_exam_hint(text: str) -> bool:
    normalized = (text or "").strip().lower()
    if not normalized:
        return False
    return any(keyword in normalized for keyword in ("toets", "pta", "tentamen", "examen"))


def _apply_vak_specific_row_postprocessing(ctx: _DocParseContext, row: DocRow) -> None:
    if not row.notities or not isinstance(row.toets, dict):
        return
    if row.toets.get("type"):
        return

    if ctx.vak.lower() == "aardrijkskunde" and _note_contains_exam_hint(row.notities):
        row.toets = {
            "type": row.notities,
            "weging": row.toets.get("weging"),
            "herkansing": row.toets.get("herkansing"),
        }
        row.notities = None


def _extract_rows_from_context(
    ctx: _DocParseContext, target_periode: Optional[int] = None
) -> List[DocRow]:
    periode = (
        target_periode
        or ctx.periode_footer
        or ctx.periode_filename
        or ctx.periode_text
    )
    schooljaar = ctx.schooljaar
    results: List[DocRow] = []

    prev_week: Optional[int] = None
    stop = False

    for table_index, (tbl, marker) in enumerate(ctx.table_markers):
        if not _table_matches_period(marker, periode):
            continue
        if stop:
            break
        rows = _table_rows_texts(tbl)
        if len(rows) < 2:
            continue
        headers, data_rows, header_row_count = _split_header_and_data_rows(rows)
        if not data_rows:
            continue
        week_col = _find_week_column(headers)
        date_col = find_header_idx(headers, DATE_HEADER_KEYWORDS)
        les_col = find_header_idx(headers, LES_HEADER_KEYWORDS)
        ond_col = find_header_idx(headers, ONDERWERP_HEADERS)
        leer_col = find_header_idx(headers, LEERDOEL_HEADERS)
        hw_col = find_header_idx(headers, HUISWERK_HEADERS)
        opd_col = find_header_idx(headers, OPDRACHT_HEADERS)
        inl_col = find_header_idx(headers, INLEVER_HEADERS)
        toets_col = find_header_idx(headers, TOETS_HEADERS)
        bron_col = find_header_idx(headers, BRON_HEADERS)
        not_col = find_header_idx(headers, NOTITIE_HEADERS)
        klas_col = find_header_idx(headers, KLAS_HEADERS)
        loc_col = find_header_idx(headers, LOCATIE_HEADERS)

        allow_wrap = (marker is not None) or (periode == 2)
        for data_offset, r in enumerate(data_rows):
            row_index = header_row_count + data_offset
            if stop:
                break
            week_text = None
            week_info = WeekCellParseResult([], None, None, None)
            if week_col is not None and week_col < len(r):
                week_text = r[week_col]
                week_info = parse_week_cell_details(week_text)
            elif date_col is not None and date_col < len(r):
                iso = parse_date_cell(r[date_col], schooljaar)
                if iso:
                    try:
                        wk = date.fromisoformat(iso).isocalendar().week
                        if 1 <= wk <= 53:
                            week_info = _make_week_result([wk], None)
                    except Exception:
                        pass
            if not week_info.weeks:
                continue

            datum = None
            datum_eind = None
            if date_col is not None and date_col < len(r):
                start_candidate, end_candidate = parse_date_range_cell(r[date_col], schooljaar)
                datum = start_candidate or datum
                if end_candidate and end_candidate != datum:
                    datum_eind = end_candidate
            if not datum and week_text:
                start_candidate, end_candidate = parse_date_range_cell(week_text, schooljaar)
                datum = start_candidate or datum
                if not datum_eind and end_candidate and end_candidate != datum:
                    datum_eind = end_candidate
            if datum_eind == datum:
                datum_eind = None

            base: Dict[str, Optional[str]] = {}
            base_list: Dict[str, Optional[List[str]]] = {}
            base_dict: Dict[str, Optional[Dict[str, Optional[str]]]] = {}

            opd_text = None
            toets_text = None
            if les_col is not None and les_col < len(r):
                base["les"] = normalize_text(r[les_col]) or None
            if ond_col is not None and ond_col < len(r):
                base["onderwerp"] = normalize_text(r[ond_col]) or None
            if leer_col is not None and leer_col < len(r):
                base_list["leerdoelen"] = split_bullets(r[leer_col])
            if hw_col is not None and hw_col < len(r):
                base["huiswerk"] = normalize_text(r[hw_col]) or None
            if opd_col is not None and opd_col < len(r):
                opd_text = r[opd_col]
                base["opdracht"] = normalize_text(opd_text) or None
            if inl_col is not None and inl_col < len(r):
                base["inleverdatum"] = parse_date_cell(r[inl_col], schooljaar)
            if toets_col is not None and toets_col < len(r):
                toets_text = r[toets_col]
                base_dict["toets"] = parse_toets_cell(toets_text)
            if bron_col is not None and bron_col < len(r):
                base_list["bronnen"] = find_urls(r[bron_col])
            if not_col is not None and not_col < len(r):
                base["notities"] = normalize_text(r[not_col]) or None
            if klas_col is not None and klas_col < len(r):
                base["klas_of_groep"] = normalize_text(r[klas_col]) or None
            if loc_col is not None and loc_col < len(r):
                base["locatie"] = normalize_text(r[loc_col]) or None

            if not base.get("inleverdatum"):
                candidate = None
                if opd_text:
                    candidate = parse_date_cell(opd_text, schooljaar)
                if not candidate and toets_text:
                    candidate = parse_date_cell(toets_text, schooljaar)
                if candidate:
                    base["inleverdatum"] = candidate

            accepted_weeks: List[int] = []
            for w in week_info.weeks:
                if not (1 <= w <= 53):
                    continue
                if _is_new_period(prev_week, w, allow_wrap=allow_wrap):
                    stop = True
                    break
                accepted_weeks.append(w)
                prev_week = w

            if stop or not accepted_weeks:
                continue

            final_week_info = _make_week_result(accepted_weeks, week_info.label)
            anchor_week = final_week_info.week_span_start
            topic_value = base.get("onderwerp") or base.get("les")
            dr = DocRow(
                week=anchor_week,
                weeks=final_week_info.weeks or None,
                week_span_start=final_week_info.week_span_start,
                week_span_end=final_week_info.week_span_end,
                week_label=final_week_info.label,
                datum=datum,
                datum_eind=datum_eind,
                les=None,
                onderwerp=topic_value,
                leerdoelen=base_list.get("leerdoelen"),
                huiswerk=base.get("huiswerk"),
                opdracht=base.get("opdracht"),
                inleverdatum=base.get("inleverdatum"),
                toets=base_dict.get("toets"),
                bronnen=base_list.get("bronnen"),
                notities=base.get("notities"),
                klas_of_groep=base.get("klas_of_groep"),
                locatie=base.get("locatie"),
                source_row_id=f"{ctx.filename}:t{table_index}:r{row_index}",
            )
            _apply_vak_specific_row_postprocessing(ctx, dr)
            results.append(dr)

    return results


def extract_rows_from_docx(
    path: str, filename: str, target_periode: Optional[int] = None
) -> List[DocRow]:
    ctx = _build_doc_context(path, filename)
    return _extract_rows_from_context(ctx, target_periode)


def extract_entries_from_docx(
    path: str, filename: str, target_periode: Optional[int] = None
) -> List[RawEntry]:
    rows = extract_rows_from_docx(path, filename, target_periode)
    return BaseParser.entries_from_rows(rows, BASE_PARSER)


def extract_all_periods_from_docx(
    path: str, filename: str
) -> List[Tuple[DocMeta, List[DocRow]]]:
    ctx = _build_doc_context(path, filename)
    periods: List[Optional[int]] = []
    seen: Set[int] = set()

    for _, marker in ctx.table_markers:
        if marker is None:
            continue
        if marker not in seen:
            periods.append(marker)
            seen.add(marker)

    for candidate in (
        ctx.periode_footer,
        ctx.periode_filename,
        ctx.periode_text,
    ):
        if candidate is None:
            continue
        if candidate not in seen:
            periods.append(candidate)
            seen.add(candidate)

    if not periods:
        periods.append(None)

    parsed: List[Tuple[DocMeta, List[DocRow]]] = []
    for periode in periods:
        meta = _extract_meta_from_context(ctx, periode)
        if not meta:
            continue
        rows = _extract_rows_from_context(ctx, periode)
        parsed.append((meta, rows))

    return parsed
