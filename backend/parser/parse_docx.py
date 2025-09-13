from typing import List, Dict
from docx import Document

def extract_items(fobj) -> List[Dict]:
    doc = Document(fobj)
    items: list[dict] = []
    # tables first
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                items.append({"date": cells[0], "title": cells[1], "text": " ".join(cells[1:])})
    # then paragraphs (fallback)
    buf_date = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        # naive detection: a line that looks like a date (dd-mm-yyyy)
        if any(ch.isdigit() for ch in t) and "-" in t and len(t) <= 12:
            buf_date = t
            continue
        items.append({"date": buf_date, "title": t, "text": t})
    return items
